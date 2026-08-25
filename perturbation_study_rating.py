"""
perturbation_study_rating.py

Runs the single-document perturbation study (profiles/1. Questionnaire_
Perturbation_Study_FollowUP.2.docx) through one or more models via
OpenRouter. Unlike rate_pairs_openrouter.py / Profiles_rating.py, this is
NOT a folder of many C##_questionnaire.docx pair files — it's one document
containing:

    - One baseline pair of profiles
    - 3 baseline evaluation questions (realism 1-7, construction a-c,
      overall compatibility 0-100)
    - 16 perturbation scenarios (S1-S16), each a single changed detail
      applied to the SAME baseline pair, each requiring one revised
      overall compatibility rating (0-100)

The scenarios live in a Word table, not paragraphs, so this script reads
both the document body and its table and assembles them into one prompt.
Nothing about the wording is hardcoded here beyond that assembly — the
actual instructions, profiles, questions, and scenario text all come
straight from the .docx.

Each model is called ONCE on this document (one row per model in the
output spreadsheet) — this is not a per-pair loop.

Requirements:
    pip install openai python-dotenv python-docx openpyxl --break-system-packages

Usage:
    Uses the same .env / OPENROUTER_API_KEY as Profiles_rating.py.
    Then just run:
        python perturbation_study_rating.py

    Edit MODELS near the top of this file to control which model(s) run.

Input:
    profiles/1. Questionnaire_Perturbation_Study_FollowUP.2.docx

Output:
    results/perturbation_followup_results.xlsx   (one row per model)
"""

import os
import re
import time
from pathlib import Path

import docx
from dotenv import load_dotenv
from openai import OpenAI, APIConnectionError, APITimeoutError, InternalServerError
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

load_dotenv()

# --- Model selection -------------------------------------------------------
# Same convention as Profiles_rating.py: any OpenRouter model slug.
# Every model listed here is called once against this one document.
MODELS = [
    "anthropic/claude-opus-5",
    "openai/gpt-5.5",
    "google/gemini-2.5-flash",
    "x-ai/grok-4.6",
    "nvidia/nemotron-3-ultra-550b-a55b",
]

TEMPERATURE = None  # left unset, same reasoning as Profiles_rating.py

MAX_TOKENS = 4000
MAX_TOKENS_CEILING = 16000

MAX_RETRIES = 4
RETRY_BACKOFF_SECONDS = 5

QUESTIONNAIRE_PATH = Path("profiles") / "1. Questionnaire_Perturbation_Study_FollowUP.2.docx"
RESULTS_DIR = Path("results")
RESULTS_DIR.mkdir(exist_ok=True)
OUT_PATH = RESULTS_DIR / "perturbation_followup_results.xlsx"

TITLE_HEADING_STYLE = "Heading 1"

# S1-S16, in the order they appear in the scenario table.
NUM_SCENARIOS = 16


def extract_prompt_text(path):
    """
    Reads the perturbation-study .docx top to bottom and returns its full
    text as a single string: all body paragraphs (skipping the "Heading
    1" section titles, which are filing labels, not task text) IN
    DOCUMENT ORDER, with the scenario table rendered as plain numbered
    lines (e.g. "S1: The man is 40 years old instead of 30; the woman
    remains 28.") inserted where the table appears in the document.

    python-docx exposes paragraphs and tables as separate lists with no
    built-in interleaving, so this walks the document body's raw XML
    children to preserve the actual top-to-bottom order rather than
    dumping all paragraphs first and the table at the very end.
    """
    d = docx.Document(path)
    body = d.element.body

    # Map each underlying XML element back to its python-docx wrapper so
    # we can pull .text (paragraphs) or iterate rows (tables) as we walk
    # the body in document order.
    paragraphs_by_id = {p._p: p for p in d.paragraphs}
    tables_by_id = {t._tbl: t for t in d.tables}

    lines = []
    for child in body.iterchildren():
        if child in paragraphs_by_id:
            p = paragraphs_by_id[child]
            text = p.text.strip()
            if not text:
                continue
            style_name = p.style.name if p.style else ""
            if style_name == TITLE_HEADING_STYLE:
                continue
            lines.append(text)
        elif child in tables_by_id:
            t = tables_by_id[child]
            for row in t.rows[1:]:  # skip the header row ("Scenario" / "Change from baseline" / ...)
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2 and cells[0]:
                    lines.append(f"{cells[0]}: {cells[1]}")

    return "\n\n".join(lines)


def rate_document(client, prompt_text, model):
    """
    Same retry/truncation-escalation logic as Profiles_rating.py's
    rate_one_pair, applied to this single document instead of a per-pair
    questionnaire.
    """
    current_max_tokens = MAX_TOKENS

    while True:
        kwargs = dict(
            model=model,
            messages=[{"role": "user", "content": prompt_text}],
            max_tokens=current_max_tokens,
        )
        if TEMPERATURE is not None:
            kwargs["temperature"] = TEMPERATURE

        last_error = None
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                response = client.chat.completions.create(**kwargs)
                choice = response.choices[0]
                truncated = choice.finish_reason == "length"

                if truncated and current_max_tokens < MAX_TOKENS_CEILING:
                    next_max_tokens = min(current_max_tokens * 2, MAX_TOKENS_CEILING)
                    print(f"  Truncated at max_tokens={current_max_tokens}, "
                          f"retrying with max_tokens={next_max_tokens} ...")
                    current_max_tokens = next_max_tokens
                    break

                if truncated:
                    print(f"  WARNING: still truncated at MAX_TOKENS_CEILING="
                          f"{MAX_TOKENS_CEILING}. Later scenarios will be blank.")

                return choice.message.content, truncated

            except (APIConnectionError, APITimeoutError, InternalServerError) as e:
                last_error = e
                if attempt < MAX_RETRIES:
                    wait = RETRY_BACKOFF_SECONDS * (2 ** (attempt - 1))
                    print(f"  Network error ({type(e).__name__}), "
                          f"retrying in {wait}s (attempt {attempt}/{MAX_RETRIES}) ...")
                    time.sleep(wait)
                else:
                    print(f"  Giving up after {MAX_RETRIES} attempts: {e}")
                    raise last_error


def parse_response(raw_text):
    """
    Parses the model's answer into the 3 baseline items plus S1-S16.

    Same marker-then-window approach as Profiles_rating.py's
    parse_numeric_fields, adapted to this document's item set: items
    1-3 are the baseline questions, and S1-S16 are the scenario
    ratings. Each marker's position bounds where its own value must be
    found (up to the next marker of any kind), so a number appearing
    inside one item's text can't leak into a neighboring item.
    """
    fields = {}

    item_marker_re = re.compile(
        r"(?:\*{0,2}\s*Item\s+(?P<n1>\d{1,2})\b(?:\s*\([^)]{0,100}\))?\s*:?\s*\*{0,2}"
        r"|(?:^|\n)\s*\**\s*(?P<n2>\d{1,2})\.(?!\d))",
        re.IGNORECASE
    )
    # Like item_marker_re, this also swallows an optional parenthetical
    # right after the marker (e.g. "**S1** (man is 40 instead of 30): 68")
    # — without that, the value window would start inside the
    # parenthetical and the first number found there (e.g. "40", the
    # restated age from the scenario description) would be wrongly
    # picked up instead of the actual rating ("68").
    scenario_marker_re = re.compile(
        r"\**\s*S(?P<s>\d{1,2})\s*\**\s*(?:\([^)]{0,200}\))?\s*:?",
        re.IGNORECASE
    )

    raw_markers = []
    for m in item_marker_re.finditer(raw_text):
        n = m.group("n1") or m.group("n2")
        if n is not None:
            raw_markers.append((("item", int(n)), m.start(), m.end()))
    for m in scenario_marker_re.finditer(raw_text):
        raw_markers.append((("scenario", int(m.group("s"))), m.start(), m.end()))

    raw_markers.sort(key=lambda x: x[1])

    seen = set()
    markers = []
    for key, start, end in raw_markers:
        if key not in seen:
            seen.add(key)
            markers.append((key, start, end))

    def value_window(key):
        matches = [(start, end) for k, start, end in markers if k == key]
        if not matches:
            return ""
        _, label_end = matches[0]
        later_starts = [start for _, start, _ in markers if start > label_end]
        window_end = min(later_starts) if later_starts else len(raw_text)
        return raw_text[label_end:window_end]

    def extract_value(key, kind):
        window = value_window(key)
        if not window:
            return None
        if kind == "q1":
            m = re.search(r"\b([1-7])\b", window)
        elif kind == "q2":
            m = re.search(r"\b([1-3]|[a-cA-C])\b", window)
        else:
            m = re.search(r"\b(\d{1,3})\b", window)
        return m.group(1) if m else None

    fields["q1_realism"] = extract_value(("item", 1), "q1")
    fields["q2_construction"] = extract_value(("item", 2), "q2")
    fields["q3_compatibility"] = extract_value(("item", 3), "q3")

    for s in range(1, NUM_SCENARIOS + 1):
        fields[f"s{s}"] = extract_value(("scenario", s), "scenario")

    # Fallback for models that answer with a bare, unlabeled list of
    # numbers (one per line, no "1.", "Item 1", or "S1" markers at all —
    # observed from grok-4.6). The marker-based parse above finds nothing
    # to anchor on in that case, so every field comes back None even
    # though the answer is right there. If literally every field is
    # still empty, and the response consists of exactly 3 + NUM_SCENARIOS
    # standalone numeric lines (optionally wrapped in bold markers),
    # assign them positionally in document order: items 1-3, then S1-S16.
    all_keys = ["q1_realism", "q2_construction", "q3_compatibility"] + [f"s{s}" for s in range(1, NUM_SCENARIOS + 1)]
    if all(fields.get(k) is None for k in all_keys):
        bare_numbers = re.findall(r"^\s*\**\s*(\d{1,3})\s*\**\s*$", raw_text, re.MULTILINE)
        if len(bare_numbers) == 3 + NUM_SCENARIOS:
            fields["q1_realism"] = bare_numbers[0]
            fields["q2_construction"] = bare_numbers[1]
            fields["q3_compatibility"] = bare_numbers[2]
            for i, s in enumerate(range(1, NUM_SCENARIOS + 1)):
                fields[f"s{s}"] = bare_numbers[3 + i]

    fields["raw_response"] = raw_text
    return fields


COLUMNS = [
    ("Model", "model"),
    ("Truncated?", "truncated"),
    ("Incomplete?", "incomplete"),
    ("Q1 Realism (1-7)", "q1_realism"),
    ("Q2 Construction (a-c)", "q2_construction"),
    ("Q3 Baseline Compatibility (0-100)", "q3_compatibility"),
] + [
    (f"S{s}", f"s{s}") for s in range(1, NUM_SCENARIOS + 1)
] + [
    ("Raw Response", "raw_response"),
]


def build_workbook(rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Perturbation Study Results"

    FONT = "Arial"
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(name=FONT, bold=True, color="FFFFFF", size=10)
    cell_font = Font(name=FONT, size=10)
    model_font = Font(name=FONT, bold=True, size=10)
    thin = Side(style="thin", color="B7B7B7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    wrap = Alignment(wrap_text=True, vertical="top")

    for c, (label, _) in enumerate(COLUMNS, start=1):
        cell = ws.cell(row=1, column=c, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[1].height = 30

    for r, row in enumerate(rows, start=2):
        for c, (_, key) in enumerate(COLUMNS, start=1):
            val = row.get(key)
            cell = ws.cell(row=r, column=c, value=val)
            cell.border = border
            if key == "model":
                cell.font = model_font
                cell.alignment = center
            elif key == "raw_response":
                cell.font = cell_font
                cell.alignment = wrap
            else:
                cell.font = cell_font
                cell.alignment = center

    for c, (_, key) in enumerate(COLUMNS, start=1):
        col_letter = get_column_letter(c)
        if key == "model":
            ws.column_dimensions[col_letter].width = 28
        elif key == "raw_response":
            ws.column_dimensions[col_letter].width = 80
        elif key in ("truncated", "incomplete"):
            ws.column_dimensions[col_letter].width = 11
        else:
            ws.column_dimensions[col_letter].width = 10

    ws.freeze_panes = "B2"

    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = "1:1"

    return wb


def main():
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        raise SystemExit("Set OPENROUTER_API_KEY (e.g. in a .env file) before running.")

    if not QUESTIONNAIRE_PATH.exists():
        raise SystemExit(f"Questionnaire not found: {QUESTIONNAIRE_PATH}")

    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
        default_headers={
            "HTTP-Referer": "https://your-site-or-project.example",
            "X-Title": "Perturbation Study Rating",
        },
    )

    prompt_text = extract_prompt_text(QUESTIONNAIRE_PATH)
    print(f"Loaded questionnaire: {QUESTIONNAIRE_PATH.name}")
    print(f"Running models: {MODELS}")

    rows = []
    for model in MODELS:
        print(f"\n=== Model: {model} ===")
        try:
            raw_text, truncated = rate_document(client, prompt_text, model=model)
            parsed = parse_response(raw_text)

            key_late_fields = ["q3_compatibility", "s1", "s16"]
            incomplete = (not truncated) and all(parsed.get(k) is None for k in key_late_fields)

            row = {
                "model": model,
                "truncated": "YES" if truncated else "",
                "incomplete": "YES" if incomplete else "",
                **parsed,
            }
        except Exception as e:
            print(f"  FAILED for {model} after retries: {e}")
            row = {"model": model, "raw_response": f"ERROR: {e}"}

        rows.append(row)

        wb = build_workbook(rows)
        wb.save(OUT_PATH)

        time.sleep(1)

    print(f"\nDone. Saved {len(rows)} rows to {OUT_PATH}")


if __name__ == "__main__":
    main()
